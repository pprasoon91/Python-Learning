import os
import docx  # Library for handling Word documents
from fpdf import FPDF  # Library for generating PDFs
import requests  # For making API requests
import json  # For handling JSON data

def generate_text(prompt):
    """Function to generate text using the Ollama (Mistral) model."""
    url = "http://localhost:11434/api/generate"  # Local server URL for Ollama API
    payload = {
        "model": "mistral",  # Specifying the model to use
        "prompt": prompt,  # Passing the prompt for text generation
        "stream": False  # Disabling streaming
    }
    headers = {"Content-Type": "application/json"}  # Setting the request headers
    
    print(f"Sending request to Ollama: {prompt[:50]}...")  # Display prompt in terminal
    
    try:
        # Sending a POST request to the Ollama API without timeout
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()  # Raises an error for HTTP issues
        generated_text = response.json().get("response", "No response received")
        
        print(f"Generated content: {generated_text[:100]}...")  # Display generated content in terminal
        
        return generated_text  # Extracting response
    except requests.exceptions.RequestException as e:
        print(f"Error generating text: {e}")  # Handling exceptions
        return "Error generating text"

def generate_book():
    """Function to generate a book in both Word and PDF formats."""
    # Create a new Word document
    doc = docx.Document()
    title = "AI in Daily Life"  # Title of the book
    doc.add_heading(title, 0)  # Adding title as heading level 0
    
    # Generate and add introduction section
    introduction = generate_text("Write an introduction about AI in daily life.")
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(introduction)
    
    # Create and add table of contents
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Smart Homes\n2. Healthcare\n3. Finance\n4. Productivity\n5. Conclusion")
    
    # Define sections with their respective prompts
    sections = [
        ("Smart Homes", "How AI-powered devices such as smart assistants and IoT devices enhance daily convenience."),
        ("Healthcare", "How AI assists in diagnostics, personalized treatment, and medical research."),
        ("Finance", "How AI helps in fraud detection, automated trading, and financial planning."),
        ("Productivity", "How AI-driven tools automate tasks, enhance efficiency, and support decision-making.")
    ]
    
    # Generate and add content for each section
    for title, prompt in sections:
        doc.add_heading(title, level=1)
        content = generate_text(prompt)
        doc.add_paragraph(content)
    
    # Generate and add conclusion
    doc.add_heading("Conclusion", level=1)
    conclusion = generate_text("Write a conclusion about the impact of AI in daily life.")
    doc.add_paragraph(conclusion)
    
    # Save the Word document
    word_filename = "AI_in_Daily_Life.docx"
    doc.save(word_filename)
    
    # Convert to PDF
    pdf = FPDF()  # Initialize PDF object
    pdf.set_auto_page_break(auto=True, margin=15)  # Enable auto page break
    pdf.add_page()  # Add a page to the PDF
    pdf.set_font("Arial", size=12)  # Set default font
    
    # Add title to PDF
    pdf.cell(200, 10, txt=title, ln=True, align='C')
    pdf.ln(10)  # Line break
    
    # Add introduction to PDF
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="Introduction", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, introduction)
    pdf.ln(5)
    
    # Add sections to PDF
    for title, prompt in sections:
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(200, 10, txt=title, ln=True)
        pdf.set_font("Arial", size=12)
        content = generate_text(prompt)
        pdf.multi_cell(0, 10, content)
        pdf.ln(5)
    
    # Add conclusion to PDF
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="Conclusion", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, conclusion)
    
    # Save the PDF file
    pdf_filename = "AI_in_Daily_Life.pdf"
    pdf.output(pdf_filename)
    
    print(f"Book saved as {word_filename} and {pdf_filename}")  # Print success message

if __name__ == "__main__":
    generate_book()  # Call function to generate the book
